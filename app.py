import streamlit as st
import pandas as pd
import pdfplumber
import json
from openai import OpenAI
import asyncio
from openai import AsyncOpenAI
import re
from docx import Document
from docx.shared import Pt
import io
import time
from streamlit_gsheets import GSheetsConnection


# ==========================================
# 💎 【现代 SaaS 风】CSS 样式注入
# ==========================================
def apply_custom_design():
    st.markdown("""
    <style>
    /* 1. 全局背景：非常淡的蓝灰色 (Slate-50)，让白色卡片能凸显出来 */
    .main {
        background-color: #F8FAFC !important;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
    }
    .block-container {
        max-width: 98% !important; /* 从 1200px 改为 98%，几乎撑满全屏 */
        padding-left: 3rem !important; /* 两侧保留基础的呼吸空间 */
        padding-right: 3rem !important;
        padding-top: 2rem !important;
        padding-bottom: 5rem !important;
    }

    /* 2. 卡片系统：白底、细腻圆角(12px)、高级柔和阴影 */
    div[data-testid="stExpander"], 
    div[data-testid="stMetric"], 
    .stDataFrame, 
    div.stTabs {
        background-color: #FFFFFF !important;
        border: 1px solid #E2E8F0 !important;
        border-radius: 12px !important;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05), 0 2px 4px -1px rgba(0, 0, 0, 0.03) !important;
        padding: 1.5rem !important;
        margin-bottom: 1.5rem !important;
    }

    /* 3. 按钮设计：使用高级的靛蓝色 (Indigo) 渐变，带有微动效 */
    .stButton>button {
        width: 100% !important;
        background: linear-gradient(135deg, #6366F1 0%, #4F46E5 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        height: 3rem !important;
        box-shadow: 0 4px 6px -1px rgba(99, 102, 241, 0.2) !important;
        transition: all 0.2s ease-in-out !important;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 15px -3px rgba(99, 102, 241, 0.3) !important;
    }

    /* 4. 标题色彩：深灰蓝，稳重且专业 */
    h1, h2, h3 {
        color: #0F172A !important;
        font-weight: 700 !important;
    }

    /* 5. 输入框美化：增加一点内阴影，显得更精致 */
    .stTextInput>div>div>input, 
    .stTextArea>div>div>textarea, 
    [data-baseweb="select"] {
        border-radius: 8px !important;
        border: 1px solid #CBD5E1 !important;
        box-shadow: inset 0 1px 2px rgba(0, 0, 0, 0.02) !important;
    }
    /* 输入框聚焦时的颜色 */
    .stTextInput>div>div>input:focus, 
    .stTextArea>div>div>textarea:focus {
        border-color: #6366F1 !important;
        box-shadow: 0 0 0 1px #6366F1 !important;
    }
    /* 修复底部聊天框过宽，强制其与主体内容等宽并居中 */
    [data-testid="stChatInput"] {
        max-width: 1400px !important; /* 限制它的最大宽度 */
        margin: 0 auto !important; /* 强制水平居中 */
    }
    
    /* 针对某些 Streamlit 版本的额外保护措施 */
    .stChatInputContainer {
        padding-left: 3rem !important; /* 与 block-container 的左右 padding 保持一致 */
        padding-right: 3rem !important;
    }
    </style>
    """, unsafe_allow_html=True)

# --- 1. 页面基本配置 ---
st.set_page_config(page_title=" 暑期实习求职利器", layout="wide", initial_sidebar_state="expanded")

# --- 现代 SaaS 风 Hero Section ---
st.markdown("""
<style>
/* 背景流体变幻动画 */
@keyframes gradientShift { 0% { background-position: 0% 50%; } 50% { background-position: 100% 50%; } 100% { background-position: 0% 50%; } }
/* 卡片上下悬浮动画 */
@keyframes floatBox { 0% { transform: translateY(0px); } 50% { transform: translateY(-8px); } 100% { transform: translateY(0px); } }

.fluid-hero-soft {
    /* 背景色保持你喜欢的流体紫 */
    background: linear-gradient(-45deg, #EEF2FF, #E0E7FF, #C7D2FE, #818CF8, #EEF2FF);
    background-size: 300% 300%;
    animation: gradientShift 10s ease infinite, floatBox 6s ease-in-out infinite;
    padding: 45px 50px;

    /* 核心修改 1：更大的圆角，削弱“矩形”感 */
    border-radius: 32px; 

    /* 核心修改 2：彻底去掉实线边框 */
    border: none; 

    /* 核心修改 3：超柔和的紫色弥散光 (外阴影) + 边缘向内渐隐的白色融合光 (内阴影) */
    box-shadow: 0 24px 50px -12px rgba(129, 140, 248, 0.25), inset 0 0 30px rgba(255, 255, 255, 0.6);
    margin-bottom: 30px;
}
</style>
""" + """<div class="fluid-hero-soft"><h1 style="color: #312E81 !important; margin-top: 0; font-size: 2.3rem; font-weight: 800; border: none;">✨ 职场之星：求职竞争力引擎</h1><p style="color: #4338CA; font-size: 1.15rem; margin-top: 15px; max-width: 700px; line-height: 1.7;">大厂实习简历总是石沉大海？让我们用 AI 深度解析你的职场基因，把平庸的描述转化为让面试官眼前一亮的“必杀技”。</p></div>""",
            unsafe_allow_html=True)

apply_custom_design()  # 调用美容函数


# 从 Secrets 获取关键配置
try:
    DEEPSEEK_API_KEY = st.secrets["DEEPSEEK_API_KEY"]
    SQL_SHEET_URL = st.secrets["SQL_SHEET_URL"]
except Exception as e:
    st.error("❌ 缺失 Secrets 配置（DEEPSEEK_API_KEY 或 SQL_SHEET_URL）")
    st.stop()

# 建立 GSheets 连接
conn = st.connection("gsheets", type=GSheetsConnection)


# --- 2. 权限与计费核心函数 ---

def verify_user(license_key):
    """验证激活码：检查是否存在、是否激活、额度是否足够"""
    try:
        # 读取用户表 (ttl=0 保证实时获取最新额度)
        user_df = conn.read(spreadsheet=SQL_SHEET_URL, worksheet="users", ttl=0)
        user_data = user_df[user_df['License_Key'] == license_key]

        if user_data.empty:
            return None, "❌ 激活码无效"

        user_info = user_data.iloc[0]
        if user_info['Status'] != 'active':
            return None, "🚫 该激活码已被禁用"

        if user_info['Used_Count'] >= user_info['Total_Count']:
            return None, "⚠️ 额度已用完，请联系管理员续费"

        return user_info, "✅ 验证通过"
    except Exception as e:
        return None, f"校验出错: {e}"


def deduct_usage(license_key, amount=1.0):
    try:
        # 1. 实时读取最新数据
        user_df = conn.read(spreadsheet=SQL_SHEET_URL, worksheet="users", ttl=0)

        # 【核心修复】强制将 Used_Count 列转为浮点数类型，防止赋值时被截断
        user_df['Used_Count'] = user_df['Used_Count'].astype(float)

        idx = user_df[user_df['License_Key'] == license_key].index[0]

        # 2. 计算新额度
        current_used = float(user_df.at[idx, 'Used_Count'])
        new_used = current_used + amount

        # 现在赋值给 float 类型的列，5.5 就不会变成 5 了
        user_df.at[idx, 'Used_Count'] = new_used

        # 3. 写回 Google Sheets
        conn.update(spreadsheet=SQL_SHEET_URL, worksheet="users", data=user_df)

        # 4. 同步更新本地缓存
        if "user_info" in st.session_state:
            st.session_state.user_info['Used_Count'] = new_used

        return True
    except Exception as e:
        st.error(f"计费系统异常: {e}")
        return False



# --- 新增：带指数退避的重试函数 ---
def call_ai_with_retry(client, model, messages, max_retries=3, delay=2):
    """
    遇到频率限制时自动重试：2s -> 4s -> 8s
    """
    for i in range(max_retries):
        try:
            return client.chat.completions.create(model=model, messages=messages)
        except Exception as e:
            # 如果是频率限制错误且还有重试机会
            if ("429" in str(e) or "rate_limit" in str(e).lower()) and i < max_retries - 1:
                wait_time = delay * (2 ** i)
                time.sleep(wait_time)
                continue
            raise e # 其他错误或重试耗尽则抛出


async def async_call_ai_with_retry(client, model, messages, max_retries=3, delay=2):
    """异步版本的重试函数"""
    for i in range(max_retries):
        try:
            return await client.chat.completions.create(model=model, messages=messages)
        except Exception as e:
            if ("429" in str(e) or "rate_limit" in str(e).lower()) and i < max_retries - 1:
                wait_time = delay * (2 ** i)
                await asyncio.sleep(wait_time)
                continue
            raise e


async def process_single_batch(batch, cv_text, batch_index, semaphore):
    """处理单个批次的异步任务"""
    async with semaphore:  # 限制最大并发数，防止被 DeepSeek 封号
        prompt = f"""
        你现在是一位拥有 15 年经验的资深招聘专家，擅长从复杂的简历中挖掘人才与岗位的深度契合点。

        ### 评估背景
        【候选人简历】：
        {cv_text[:2500]}

        【待匹配岗位列表】：
        {json.dumps(batch, ensure_ascii=False)}

        ### 你的任务
        请基于以下逻辑框架，对简历与每个岗位进行深度匹配分析：

        1. **核心技能匹配度**：对比简历中的技术栈（如 Python, SQL, 财务建模等）与 JD 的硬性要求。
        2. **行业/项目相关性**：分析过往项目或实习经历在业务逻辑上是否与目标岗位一致。
        3. **软实力与潜力**：从奖项、社团经历中评估候选人的学习能力和执行力。

        ### 评分准则
        - **90-100分**：完美匹配，几乎无需培训即可上手。
        - **70-89分**：具备核心能力，但在特定经验或次要工具上略有欠缺。
        - **50-69分**：有一定基础，但需要大量带教或转岗跨度较大。
        - **50分以下**：基本不匹配。

        ### 输出要求
        请严格按 JSON 数组格式返回，不要包含任何前导语或总结语。格式如下：
        [
          {{
            "index": 岗位索引号,
            "match_score": 整数评分,
            "match_reason": "【核心优势】：[列出1-2点最匹配的经历或技能]；【潜在挑战】：[指出简历中缺少的关键要素或不足]；【综合判定】：[一句话说明为什么值得投递]。"
          }}
        ]
        """

        try:
            response = await async_call_ai_with_retry(
                aclient,
                "deepseek-chat",
                [{"role": "user", "content": prompt}]
            )

            raw_content = response.choices[0].message.content.strip()
            if raw_content.startswith("```json"):
                raw_content = raw_content.replace("```json", "").replace("```", "").strip()
            elif raw_content.startswith("```"):
                raw_content = raw_content.replace("```", "").strip()

            ai_res = json.loads(raw_content)

            # 标准化返回格式
            if isinstance(ai_res, list):
                return ai_res, batch_index, None
            elif isinstance(ai_res, dict):
                return ai_res.get("results", ai_res.get("matches", list(ai_res.values())[0])), batch_index, None

        except Exception as e:
            return None, batch_index, str(e)

# 将结果导出到word中
def export_to_word(summary, analysis, refined_data):
    doc = Document()

    # 设置全局字体（可选）
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)

    # 1. 标题
    doc.add_heading('简历深度优化报告', 0)

    # 2. 整体求职策略
    doc.add_heading('一、首席人才官：整体求职策略建议', level=1)
    doc.add_paragraph(summary)

    # 3. 岗位胜任力解析
    doc.add_heading('二、岗位胜任力深度画像', level=1)
    doc.add_paragraph(analysis)

    # 4. 各模块精修建议
    doc.add_heading('三、简历各模块精修建议', level=1)

    for section_name, content in refined_data.items():
        doc.add_heading(f'模块：{section_name}', level=2)
        # AI 返回的是 Markdown 格式，这里简单处理：
        # 如果你想要更完美的表格导出，需要解析 Markdown 表格，
        # 这里先以文本流形式导出，保证内容完整和分段清晰。
        doc.add_paragraph(content)
        doc.add_page_break()  # 每个大模块换一页，保持整洁

    # 将文档保存到内存流
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def split_resume_by_sections(text):
    """
    使用正则表达式将简历文本切分为模块
    """
    # 定义常见的简历标题关键词
    patterns = {
        "工作经历": r"(工作经历|实习|实习经历|Experience|Work)",
        "项目经历": r"(项目经历|项目经验|个人项目|Projects)",
        "技能证书": r"(专业技能|技能特长|语言能力|证书|Skills)",
        "自我评价": r"(自我评价|个人总结|About Me)"
    }

    # 1. 寻找所有可能的切分点
    matches = []
    for section, pattern in patterns.items():
        # 使用 re.MULTILINE 尝试匹配行首，减少正文干扰
        for match in re.finditer(pattern, text, re.IGNORECASE):
            matches.append((match.start(), section))

    # 2. 按文本位置排序
    matches.sort()

    # 3. 提取各个模块
    sections = {}

    # 提取“基本信息”（第一个标题之前的内容）
    if matches and matches[0][0] > 0:
        sections["基本信息"] = text[:matches[0][0]].strip()

    for i in range(len(matches)):
        start_pos, section_name = matches[i]
        # 结尾是下一个匹配点的开始，或者是全文结尾
        end_pos = matches[i + 1][0] if i + 1 < len(matches) else len(text)

        content = text[start_pos:end_pos].strip()

        # 避免内容重复：如果一个标题被匹配了多次，合并它们
        if section_name in sections:
            sections[section_name] += "\n" + content
        else:
            sections[section_name] = content

    # 4. 兜底处理：如果没有匹配到任何标题
    if not sections:
        return {"完整简历": text}

    return sections


# --- 2. 隐私保护声明 ---
with st.expander("🛡️ 隐私保护与数据安全说明", expanded=False):
    st.info("""
    **本工具郑重承诺：**
    1. **不留痕迹**：你上传的简历和岗位表仅在服务器内存中实时处理。
    2. **不存储文件**：一旦你刷新或关闭网页，所有数据将自动彻底销毁。
    3. **脱敏建议**：你可以删除简历中的手机号、住址等敏感信息，不影响 AI 评估。
    """)

# --- 3. 侧边栏：授权管理 ---

with st.sidebar:
    st.header("🔑 访问授权")
    user_code = st.text_input("请输入您的专属激活码", type="password", help="联系管理员获取")

    if not user_code:
        # 如果清空了输入框，也重置验证状态
        if "user_info" in st.session_state:
            del st.session_state.user_info
        st.info("💡 请输入激活码以解锁简历优化功能。")
        st.stop()

    # 只有在以下情况才去读取 Google Sheets:
    # 1. session_state 里没有用户信息
    # 2. 用户输入的 code 和上次验证成功的 code 不一致
    if "user_info" not in st.session_state or st.session_state.get("last_verified_code") != user_code:

        with st.spinner("正在验证权限..."):
            user_data, msg = verify_user(user_code)

            if user_data is not None:
                # 验证成功，存入“记忆”
                st.session_state.user_info = user_data.to_dict()  # 转为字典方便存储
                st.session_state.last_verified_code = user_code
            else:
                # 验证失败，显示错误并停止
                st.error(msg)
                if "user_info" in st.session_state:
                    del st.session_state.user_info
                st.stop()

    # 从“记忆”中直接读取用户信息，不再请求网络
    current_user = st.session_state.user_info

    st.success(f"欢迎回来，{current_user['User_Name']}！")

    # 计算剩余额度（注意：扣费后需要手动更新这个显示）
    remaining = current_user['Total_Count'] - current_user['Used_Count']
    st.metric("剩余可用额度", f"{remaining} 次")

    st.divider()
    st.header("🎨 简历定制偏好")
    opt_style = st.radio("文风倾向", ["稳重务实型", "极简干练型", "充满活力型"], index=1)
    detail_depth = st.select_slider("细节挖掘深度", options=["点到为止", "标准修饰", "深度重构"], value="标准修饰")

# 初始化 AI 客户端 (统一使用 Secret)
client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
aclient = AsyncOpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

# --- 5. 功能区 ---
st.header("01 / 岗位精准匹配")
# 功能一：精准匹配
# --- 修改后 ---
try:
    df = conn.read(spreadsheet=SQL_SHEET_URL, worksheet="jobs", ttl=600)
    # 变成一行浅灰色的细小提示，并展示岗位总数，增加产品的“厚重感”
    st.caption(f"🟢 岗位数据中心已就绪 | 实时在库岗位：{len(df)} 个")
except:
    st.error("无法同步岗位库，请检查表格名称是否为 'jobs'")
    st.stop()

cv_file = st.file_uploader("上传你的简历 (PDF)", type=["pdf"])
st.info("⚠️ 温馨提示：为了保证 AI 提取准确率，请务必上传由 Word 或纯文本直接导出的 PDF。暂不支持由图片转换的扫描件哦！")

if cv_file:

    # 筛选 UI 界面
    # --- 修改后 ---
    st.markdown("#### 🔍 岗位筛选")  # 用 markdown 的 h4 显得比 subheader 更精致一点

    # 🌟 核心修改：用带有 border 的容器把所有筛选框包裹起来
    with st.container(border=True):
        row1_col1, row1_col2, row1_col3 = st.columns(3)
        with row1_col1:
            category_list = df['领域大类'].dropna().unique().tolist() if '领域大类' in df.columns else []
            sel_category = st.multiselect("领域大类", options=category_list)
        with row1_col2:
            field_list = df['实习领域'].dropna().unique().tolist() if '实习领域' in df.columns else []
            sel_field = st.multiselect("实习领域", options=field_list)
        with row1_col3:
            city_list = df['工作地点'].dropna().unique().tolist() if '工作地点' in df.columns else []
            sel_cities = st.multiselect("工作地点", options=city_list)

    # 【修改】第二排筛选条件：实习时长、转正机会、学历要求
        row2_col1, row2_col2, row2_col3 = st.columns(3)
        with row2_col1:
            month_list = df['实习月数'].dropna().unique().tolist() if '实习月数' in df.columns else []
            sel_months = st.multiselect("实习时长 (月数)", options=month_list)
        with row2_col2:
            convert_list = df['转正机会'].dropna().unique().tolist() if '转正机会' in df.columns else []
            sel_convert = st.multiselect("转正机会", options=convert_list)
        with row2_col3:
            edu_list = df['学历要求'].dropna().unique().tolist() if '学历要求' in df.columns else []
            sel_edu = st.multiselect("学历要求", options=edu_list)

    # 执行 Python 过滤逻辑
    filtered_df = df.copy()
    if sel_category:
        filtered_df = filtered_df[filtered_df['领域大类'].isin(sel_category)]
    if sel_field:
        filtered_df = filtered_df[filtered_df['实习领域'].isin(sel_field)]
    if sel_cities:
        filtered_df = filtered_df[filtered_df['工作地点'].isin(sel_cities)]
    if sel_months:
        filtered_df = filtered_df[filtered_df['实习月数'].isin(sel_months)]
    if sel_convert:
        filtered_df = filtered_df[filtered_df['转正机会'].isin(sel_convert)]
    # 【新增】学历要求过滤
    if sel_edu:
        filtered_df = filtered_df[filtered_df['学历要求'].isin(sel_edu)]
        # 强制在后端按“发布时间”降序排列，确保后续预览和喂给 AI 的都是最新岗位
    if '发布时间' in filtered_df.columns:
        filtered_df = filtered_df.sort_values(by='发布时间', ascending=False)

    st.write(f"📊 筛选后符合要求的岗位：**{len(filtered_df)}** 个")
    st.dataframe(filtered_df.head(50), use_container_width=True)  # 预览前50条

    if st.button("🔥 开始 AI 智能匹配(消耗3额度，耗时3-5分钟)"):
        if filtered_df.empty:
            st.error("筛选后没有符合要求的岗位，请放宽筛选条件。")
        else:
            # 设定要交给 AI 深度评估的最大岗位数量（避免费用过高和等待太久，建议截取前 60 个）
            max_eval_count = 60
            jobs_to_eval = filtered_df.head(max_eval_count)

            with st.status(f"🚀 AI 正在深度解析前 {len(jobs_to_eval)} 个优质岗位...", expanded=True) as status:
                # 读取简历内容
                with pdfplumber.open(cv_file) as pdf:
                    cv_text = "".join([page.extract_text() for page in pdf.pages])

                # 提取关键信息
                jobs_list = jobs_to_eval[['职位名称', '职位描述', '任职要求']].reset_index().to_dict(orient='records')

                batch_size = 12
                total_batches = (len(jobs_list) + batch_size - 1) // batch_size
                progress_bar = st.progress(0)

                status.write(f"⏳ 正在将 {len(jobs_list)} 个岗位拆分为 {total_batches} 批次，启动并发解析...")


                # 定义异步执行的主控中心
                async def run_concurrent_batches():
                    # 并发锁：最多同时向 DeepSeek 发起 3 个请求（这很关键，防止报错）
                    semaphore = asyncio.Semaphore(3)
                    tasks = []

                    # 把所有的批次任务都打包好
                    for i in range(0, len(jobs_list), batch_size):
                        batch = jobs_list[i:i + batch_size]
                        batch_index = i // batch_size
                        tasks.append(process_single_batch(batch, cv_text, batch_index, semaphore))

                    all_match_data_temp = []
                    completed_count = 0

                    # 谁先完成，就先处理谁的数据，并更新进度条
                    for future in asyncio.as_completed(tasks):
                        res, b_idx, error = await future
                        completed_count += 1

                        if error:
                            status.write(f"⚠️ 第 {b_idx + 1} 批解析出现小波动，已跳过。错误：{error}")
                        elif res:
                            all_match_data_temp.extend(res)
                            status.write(f"✅ 第 {b_idx + 1} 批解析完成！")

                        progress_bar.progress(completed_count / total_batches)

                    return all_match_data_temp


                # 正式启动并发！
                all_match_data = asyncio.run(run_concurrent_batches())

                # --- 循环结束，处理所有结果 ---
                if not all_match_data:
                    st.error("所有批次的 AI 解析均失败，请稍后重试。")
                    st.stop()

                status.write("✨ 所有岗位解析完毕，正在生成最终报告...")

                # 转换为 DataFrame 并与原表合并
                ai_df = pd.DataFrame(all_match_data)

                if 'index' in ai_df.columns:
                    ai_df['index'] = ai_df['index'].astype(int)
                    final_df = filtered_df.reset_index().merge(ai_df, on='index', how='inner')

                    # 【核心修改】：在这里把 AI 返回的英文列名改成中文
                    final_df = final_df.rename(columns={'match_score': '匹配分数', 'match_reason': '匹配依据'})

                    # 将改名后的这两列排在最前面，并按“匹配分数”降序排列
                    cols = ['匹配分数', '匹配依据'] + [c for c in final_df.columns if
                                                       c not in ['匹配分数', '匹配依据', 'index']]
                    final_df = final_df[cols].sort_values(by='匹配分数', ascending=False)

                    if deduct_usage(user_code, amount=3.0):
                        pass

                    st.session_state.match_results = final_df
                    status.update(label="✅ 匹配完成！已按匹配度降序排列(本次消耗 3 次额度)", state="complete",
                                  expanded=False)
                else:
                    st.error("AI 返回的数据缺少 index 字段合并失败。")

    if "match_results" in st.session_state:
        st.subheader("🎯 匹配结果推送 (含全字段信息)")
        # 【新增需求】：匹配分数定义说明
        st.info("""
        **💡 匹配分数定义：**
        - **90-100分**：完美匹配，几乎无需培训即可上手。
        - **70-89分**：具备核心能力，但在特定经验或次要工具上略有欠缺。
        - **50-69分**：有一定基础，但需要大量带教或转岗跨度较大。
        - **50分以下**：基本不匹配。
        """)
        # 从“记忆”中读取表格并展示
        st.dataframe(st.session_state.match_results, use_container_width=True)

        # 下载按钮（现在它可以被无限次点击了）
        csv_data = st.session_state.match_results.to_csv(index=False).encode('utf-8-sig')
        st.download_button("📥 下载完整分析报告 (CSV)", data=csv_data, file_name="实习匹配结果.csv")

# --- 5. 功能二：简历深度优化 (全量精修版 - 改进手动粘贴功能) ---
st.divider()

# 初始化 Session State 缓存结果
if "refined_results" not in st.session_state:
    st.session_state.refined_results = None

# 【核心优化 1：渐进式隐藏】只有当 cv_file 存在（用户上传了简历）时，才显示第二步！
if cv_file is not None:
    st.header("02 / 简历深度优化工作台")

    # 【核心优化 2：左右分栏工作台】将页面分为左右 1:1 的两栏
    col_left, col_right = st.columns([1, 1], gap="large")

    final_sections = {}  # 用于存储最终传递给 AI 的内容

    # ================= 左侧栏：展示和提取简历内容 =================
    with col_left:
        st.markdown("#### 📄 你的简历内容")
        input_tab1, input_tab2 = st.tabs(["自动从 PDF 提取", "手动粘贴/修正"])

        with input_tab1:
            # 因为最外层已经判断了 cv_file 存在，所以这里直接解析
            with pdfplumber.open(cv_file) as pdf:
                cv_raw_text = "".join([page.extract_text() for page in pdf.pages])

            auto_sections = split_resume_by_sections(cv_raw_text)
            st.success("✅ 已从 PDF 自动识别模块，可在此微调：")

            for sec_name, sec_content in auto_sections.items():
                final_sections[sec_name] = st.text_area(
                    sec_name,
                    value=sec_content,
                    height=150,
                    key=f"auto_{sec_name}_{cv_file.name}"  # ✨ 动态身份证号
                )

        with input_tab2:
            st.caption("如果自动提取不准，可在此手动覆盖粘贴：")
            manual_sections = {
                "基本信息": st.text_area("基本信息", placeholder="例如：张三，电话，邮箱，武汉大学...", height=100),
                "工作经历": st.text_area("工作经历", placeholder="例如：2022.01-2023.01 XX公司 实习生\n1. 负责...",
                                         height=150),
                "项目经历": st.text_area("项目经历", placeholder="例如：XX数据分析项目\n使用Python进行...", height=150),
                "技能证书": st.text_area("技能证书", placeholder="例如：英语六级、Python熟练、CPA...", height=100)
            }
            # 【核心修改】：智能合并逻辑
            for k, v in manual_sections.items():
                if v.strip():  # 判断：如果用户在这个框里输入了有效文字
                    final_sections[k] = v  # 就把这个模块的内容存进 final_sections，覆盖或者新增

        # ================= 右侧栏：输入目标 JD =================
    with col_right:
        st.markdown("#### 🎯 目标岗位 JD")
        target_jd = st.text_area(
            "请贴入目标岗位要求 (JD)",
            height=580,
            placeholder="请在此粘贴完整的任职要求和职位描述...",
            label_visibility="collapsed"
        )
        st.markdown("<br>", unsafe_allow_html=True)



        # 仅仅在这里获取按钮的点击状态，但不在这里执行逻辑！
        start_btn_clicked = st.button("🪄 启动专家级精修（消耗3额度，耗时3-5分钟）", use_container_width=True)

        # ================= 🚀 全局全宽进度区 (放在左右分栏之外) =================
        # 【修改点 2】：把处理逻辑移到 col_left 和 col_right 的外面，这样它就会撑满全屏
    if start_btn_clicked:
        if not final_sections or not target_jd:
            st.error("⚠️ 请确保已确认简历内容并粘贴了目标 JD！")
        else:
            refined_data = {}
            competency_analysis = "分析生成失败"
            final_summary = "总结生成失败"

            # 加一点上边距，让加载框不要紧贴着上面的输入框
            st.markdown("<br>", unsafe_allow_html=True)
            # =================🚀 全局全宽进度区优化 =================
            # 【核心修改点】：创建一个容器，该容器会自然撑满全局 block-container 的宽度
            wide_status_container = st.container()
            with wide_status_container:
                # 将 st.status(...) 放在容器内。它就会自动左边对齐 col_left，右边对齐 col_right 了！
                with st.status("🚀 专家正在深度重构中...", expanded=True) as status:
                    total_steps = 1 + len(final_sections) + 1
                    current_step = 0
                    progress_bar = st.progress(0)

                    # --- 第一阶段：岗位胜任力解析 ---
                    status.write("🕵️ 正在进行岗位胜任力深度画像...")
                    analysis_prompt = f"""
                        你现在是资深猎头专家。请针对目标岗位 JD 进行深度解析：
                        目标岗位: {target_jd}
                        任务：提炼出企业最看重的【三项核心能力】，按以下格式输出：
                        ### 【岗位胜任力解析】
                        1. **专业能力 (Hard Skills)**：内容...
                        2. **通用素质 (Soft Skills)**：内容...
                        3. **业务潜力 (Potential)**：内容...
                        """
                    try:
                        analysis_res = call_ai_with_retry(client, "deepseek-chat",
                                                          [{"role": "user", "content": analysis_prompt}])
                        competency_analysis = analysis_res.choices[0].message.content
                    except Exception as e:
                        status.write(f"⚠️ 解析岗位失败: {e}")

                    # 更新进度条
                    current_step += 1
                    progress_bar.progress(current_step / total_steps)

                    # --- 第二阶段：逐个模块重构 ---
                    for section_name, section_content in final_sections.items():
                        if not section_content.strip(): continue
                        status.write(f"正在重构：{section_name}...")

                        # 💡 注意：这里贴回你原本那个非常详细的 specific_prompt
                        # 工业级 Prompt 2.0：高稳定性、防幻觉、去废话
                        # 工业级 Prompt 2.0：高稳定性、防幻觉、去废话 (加入 One-Shot 示例护航)
                        specific_prompt = f"""
                        你是一位拥有 15 年一线大厂招聘经验的【资深职业导师】。你的任务是针对候选人的具体经历，结合目标岗位需求，进行像素级的简历优化。

                        <INPUT_DATA>
                        <TARGET_JD>
                        {target_jd}
                        </TARGET_JD>

                        <CURRENT_SECTION_NAME>
                        {section_name}
                        </CURRENT_SECTION_NAME>

                        <CURRENT_SECTION_CONTENT>
                        {section_content}
                        </CURRENT_SECTION_CONTENT>

                        <STYLE_PREFERENCE>
                        {opt_style} (重构深度: {detail_depth})
                        </STYLE_PREFERENCE>
                        </INPUT_DATA>

                        <CRITICAL_RULES>
                        1. **绝对客观（零幻觉）**：严禁虚构任何数据、公司名称或未提及的技术栈。遇到缺失的具体数据，必须使用 `[XX]` 作为占位符！
                        2. **禁止废话（Zero-Chatter）**：直接输出最终的 Markdown 结果。绝不允许出现“好的”、“这是为您优化的结果”等客套话！
                        3. **严禁阉割工作细节（最高红线）**：原简历中的**所有具体动作描述必须保留并在“优化建议”列中扩写**！绝对不允许只提取公司名称或头衔而把具体干的活删掉！
                        4. **Markdown 表格安全（致命红线）**：表格单元格内**绝对禁止使用回车键（\\n）**！无论是【原始描述】还是【优化建议】，只要需要换行或分点，**必须且只能使用 `<br>` 标签**（例：`1. 第一点<br>2. 第二点`）。
                        </CRITICAL_RULES>
                        5.格式禁令：严禁输出任何形式的删除线（包括 ~~、<s>、<strike>）。
                        6.1:1 镜像存证：左侧“原始描述”列仅允许进行“换行符转 <br>”的操作，严禁对文字内容进行任何增删改，严禁标注哪些是“旧的”。

                        <OUTPUT_FORMAT_INSTRUCTIONS>
                        请根据 <CURRENT_SECTION_NAME> 的内容属性，选择唯一对应的输出路径：

                        ▶ **路径 A：属于“工作/实习/项目/校园经历”等包含具体动作的模块**
                        严格输出一个三列表格，以及一个追问模块。

                        🚨 **请务必严格参考以下【示例】的颗粒度和格式进行输出：**

                        ### 【标准输出示例示范】
                        #### 🛠️ 简历精修对比表
                        | 原始描述 | 优化建议 (必须包含 [XX] 占位符引导补充数据) | 优化逻辑 |
                        | :--- | :--- | :--- |
                        | 2024.02-02 XX事务所 实习生<br>• 填写底稿，运用函数核对资料<br>• 盘点现金 | **1. 业务执行：**独立负责 [XX] 家企业的审计底稿编制，运用 VLOOKUP/SUMIF 等函数处理 [XX] 万条财务数据，提升核对效率 [XX]%。<br>**2. 跨区协同：**跨部门对接 [XX] 位财务人员，完成核心资料核验。<br>**3. 资产清查：**规范执行库存现金盘点，排查并解决 [XX] 笔账目差异。 | 将日常动作转化为“业务成果+数据量化”，突显 Excel 数据处理能力，完美契合 JD 中的“财务分析核对”要求。 |

                        ### 【请按以下格式输出你的实际生成结果】
                        #### 🛠️ 简历精修对比表
                        | 原始描述 | 优化建议 (必须包含 [XX] 占位符引导补充数据) | 优化逻辑 |
                        | :--- | :--- | :--- |
                        | (将原内容填入，此处仅允许复制，原本所有的换行必须替换为 <br> 标签) | (使用 XYZ 公式深度重写每一条经历，换行必须用 <br>) | (解释这样改如何契合 JD 要求) |

                        #### 🔍 深度溯源追问 (引导候选人填补 [XX])
                        (列出 3-5 个具体、尖锐的问题，针对性地引导候选人回忆能支撑 JD 的具体数据、规模或动作细节。)
                        ---

                        ▶ **路径 B：属于“基本信息/教育背景/技能证书”等纯客观列表模块**
                        此类信息属于客观事实，严禁过度包装。严格输出一个两列表格：

                        #### 🛠️ 信息规范审查表
                        | 原始描述 | 专家备注 |
                        | :--- | :--- |
                        | (原封不动复制原句，换行替换为 <br>) | (例如：“客观学历信息已保留”、“建议将技能按熟练度分类”等) |
                        </OUTPUT_FORMAT_INSTRUCTIONS>

                        请立即开始执行，只输出最终的 Markdown 代码：
                        """

                        try:
                            module_res = call_ai_with_retry(client, "deepseek-chat",
                                                            [{"role": "user", "content": specific_prompt}])
                            raw_output = module_res.choices[0].message.content

                            # 【核心修正】：升级正则，拦截 ~~, <s>, <strike>, <del> 以及单波浪线 ~
                            # 增加对 HTML 标签的拦截，因为你开启了 unsafe_allow_html=True
                            clean_output = re.sub(r'~~|~|<s>|</s>|<strike>|</strike>|<del>|</del>', '', raw_output)

                            refined_data[section_name] = clean_output
                        except Exception as e:
                            status.write(f"⚠️ 重构 {section_name} 失败: {e}")

                        # 更新进度条
                        current_step += 1
                        progress_bar.progress(current_step / total_steps)

                    # --- 第三阶段：生成全局总结 ---
                    if refined_data:
                        status.write("📝 正在生成全局求职策略建议...")
                        all_refined = "\n".join(list(refined_data.values()))
                        summary_prompt = f"针对以下精修后的内容，总结核心竞争力、面试建议并写一段100字自我评价：\n{all_refined[:2000]}"
                        try:
                            summary_res = call_ai_with_retry(client, "deepseek-chat",
                                                             [{"role": "user", "content": summary_prompt}])
                            final_summary = summary_res.choices[0].message.content
                        except Exception as e:
                            final_summary = f"总结生成失败，错误原因：{e}"

                    # 最后将进度条拉满至 100%
                    progress_bar.progress(1.0)

                    if deduct_usage(user_code, amount=3.0):
                        pass
                    status.update(label="✅ 全量精修完成！（本次消耗3次额度）", state="complete", expanded=False)

                st.session_state.refined_results = {
                    "refined_data": refined_data,
                    "competency_analysis": competency_analysis,
                    "final_summary": final_summary
                }
                st.rerun()

    # =========================================================
    # 下方紧接着你的： # --- 结果展示与导出区 ---

 # --- 结果展示与导出区 ---
if "refined_results" in st.session_state and st.session_state.refined_results:
    results = st.session_state.refined_results

    st.divider()  # 视觉分割线
    st.success("✨ 优化方案已生成！")

    # 1. 导出按钮
    col_dl, _ = st.columns([1, 2])
    with col_dl:
        st.download_button(
            "📥 一键导出 Word 报告",
            data=export_to_word(results["final_summary"], results["competency_analysis"], results["refined_data"]),
            file_name="简历深度优化报告.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # 2. 展示全局总结和胜任力分析
    with st.expander("🎯 首席人才官：整体求职策略建议", expanded=True):
        st.markdown(results["final_summary"])

    st.markdown(results["competency_analysis"])

    # 3. 展示各模块 Tab
    st.subheader("📝 简历各模块精修建议")
    tabs = st.tabs(list(results["refined_data"].keys()))
    for i, (name, content) in enumerate(results["refined_data"].items()):
        with tabs[i]:
            st.markdown(content, unsafe_allow_html=True)

    st.info("💡 **小贴士**：优化建议中的 **[XX]** 是 AI 为你预留的数据位")

# --- 6. 功能三：交互式 AI 助手 ---
# 🌟 核心修改：加了这个判断，只有上传简历后才显示 03 模块
if cv_file is not None:
    st.divider()
    st.header("03 / 简历精修对话室")
    # 温馨提示
    st.info("💡 **计费说明**：对话模式每次提问消耗 **1** 次额度。")

    if "messages" not in st.session_state:
        st.session_state.messages = []

    # 显示聊天历史
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # 聊天输入框
    if chat_input := st.chat_input("针对优化结果，你可以继续追问"):

        # 1. 检查余额 (保持不变)
        current_u = st.session_state.get("user_info")
        if current_u['Used_Count'] >= current_u['Total_Count']:
            st.warning("⚠️ 您的额度已耗尽，请联系管理员续费后再对话。")
            st.stop()

        # 2. 【核心修改】构建上下文背景
        # 尝试从之前的步骤中抓取数据
        context_info = ""

        # 获取 JD
        current_jd = target_jd if 'target_jd' in locals() else "未提供"

        # 获取之前的精修建议
        refined_summary = ""
        if "refined_results" in st.session_state and st.session_state.refined_results:
            res = st.session_state.refined_results
            refined_summary = f"你之前给出的优化策略是：{res['final_summary']}\n"
            # 也可以把各模块的精修点简要带入
            for sec, content in res['refined_data'].items():
                refined_summary += f"--- {sec} 模块优化建议 ---\n{content[:500]}...\n"

        # 构建一个强大的 System Message
        system_prompt = f"""你是一个资深简历优化专家。
你正在协助用户进行简历修饰。以下是当前任务的背景信息，请务必基于这些信息回答用户的追问：

【目标岗位 JD】：
{current_jd}

【之前的优化成果】：
{refined_summary}

请根据以上背景，结合用户的具体提问，给出针对性、专业且简洁的修改建议。"""

        # 3. 展示并记录用户消息
        st.session_state.messages.append({"role": "user", "content": chat_input})
        with st.chat_message("user"):
            st.markdown(chat_input)

        # 4. AI 响应
        with st.chat_message("assistant"):
            with st.spinner("专家正在思考中..."):
                try:
                    # 【核心修改】将 system_prompt 作为第一条消息发送
                    full_messages = [{"role": "system", "content": system_prompt}] + st.session_state.messages

                    response = call_ai_with_retry(
                        client,
                        "deepseek-chat",
                        full_messages
                    )
                    ans = response.choices[0].message.content
                    st.markdown(ans)

                    # 5. 成功后执行扣费
                    if deduct_usage(user_code, amount=1.0):
                        st.toast(f"已消耗 1 次额度", icon="💰")

                    st.session_state.messages.append({"role": "assistant", "content": ans})

                except Exception as e:
                    st.error(f"对话中断，请重试。错误信息：{e}")

